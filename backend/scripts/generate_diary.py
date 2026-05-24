
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_diary():
    doc = Document()

    # Title
    title = doc.add_heading('NHẬT KÝ PHÁT TRIỂN DỰ ÁN MEDBOOK - 20 NGÀY HOÀN THIỆN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    p = doc.add_paragraph()
    p.add_run('Người thực hiện: ').bold = True
    p.add_run('Trương Thế Hải Thịnh & Nguyễn Thị Quỳnh Trang\n')
    p.add_run('Thời gian: ').bold = True
    p.add_run('20 Ngày (Hành trình từ Ý tưởng đến MVP)')

    # Giai đoạn 1
    doc.add_heading('GIAI ĐOẠN 1: THIẾT LẬP NỀN TẢNG (NGÀY 1 - 5)', level=1)
    p = doc.add_paragraph()
    p.add_run('Công việc: ').bold = True
    p.add_run('Thiết lập cấu trúc Backend FastAPI, kết nối PostgreSQL và xây dựng hệ thống xác thực JWT cho 3 vai trò.\n')
    p.add_run('Khó khăn: ').bold = True
    p.add_run('Phân quyền roles (Patient, Doctor, Admin) ngay từ đầu khá phức tạp. Đảm bảo bác sĩ phải được duyệt mới được hoạt động.\n')
    p.add_run('Thay đổi: ').bold = True
    p.add_run('Chuyển từ SQLite sang PostgreSQL để sử dụng tính năng SELECT FOR UPDATE chống trùng lịch.')

    # Giai đoạn 2
    doc.add_heading('GIAI ĐOẠN 2: THIẾT KẾ BỘ MÁY LỊCH HẸN (NGÀY 6 - 10)', level=1)
    p = doc.add_paragraph()
    p.add_run('Công việc: ').bold = True
    p.add_run('Phát triển Smart Scheduling Engine. Tự động tính toán slot trống theo thời gian thực dựa trên lịch mẫu của bác sĩ.\n')
    p.add_run('Khó khăn: ').bold = True
    p.add_run('Xử lý các tình huống giờ giấc lẻ phút và validate lịch biểu không chồng chéo nhau.\n')
    p.add_run('Thay đổi: ').bold = True
    p.add_run('Cho phép bác sĩ tạo "Lịch mẫu tuần" thay vì phải nhập từng ngày thủ công, giúp trải nghiệm mượt mà hơn.')

    # Giai đoạn 3: Luồng Nghiệp vụ & Thông báo
    doc.add_heading('GIAI ĐOẠN 3: LUỒNG NGHIỆP VỤ & EMAIL (NGÀY 11 - 15)', level=1)
    p = doc.add_paragraph()
    p.add_run('Công việc: ').bold = True
    p.add_run('Hoàn thiện luồng đặt lịch (Pending -> Confirm -> Completed). Tích hợp gửi email tự động xác nhận lịch hẹn.\n')
    p.add_run('Khó khăn: ').bold = True
    p.add_run('Gửi email làm chậm tốc độ phản hồi API. Đã phải chuyển sang dùng BackgroundTasks để xử lý ẩn.\n')
    p.add_run('Kỹ thuật: ').bold = True
    p.add_run('Áp dụng triệt để State Machine để quản lý trạng thái lịch hẹn, không cho phép ca đã hủy được quay lại trạng thái chờ.')

    # Giai đoạn 4: Đơn thuốc & In ấn
    doc.add_heading('GIAI ĐOẠN 4: CHUYÊN NGHIỆP HÓA & PRINTING (NGÀY 16 - 20)', level=1)
    p = doc.add_paragraph()
    p.add_run('Công việc: ').bold = True
    p.add_run('Xây dựng module kê đơn thuốc, thiết kế bản in A5 và bộ nhận diện chuyên khoa qua Icon.\n')
    p.add_run('Khó khăn: ').bold = True
    p.add_run('Tinh chỉnh CSS Print Media để bản in A5 khớp tuyệt đối trên giấy, không bị lệch hàng hay tràn trang.\n')
    p.add_run('Nâng cấp quan trọng: ').bold = True
    p.add_run('Thêm tính năng "Chỉnh sửa ca khám" vì thực tế bác sĩ có thể ghi nhầm khi chẩn đoán. Đây là thay đổi lớn giúp hệ thống thực tế hơn.')

    # Kết luận
    doc.add_heading('KẾT LUẬN', level=1)
    doc.add_paragraph("Hành trình 20 ngày đã biến MedBook từ bản vẽ proposal thành một sản phẩm MVP hoàn chỉnh, có tính ứng dụng cao và giao diện chuyên nghiệp. Mọi khó khăn kỹ thuật đều đã được giải quyết bằng các giải pháp kiến trúc bền vững.")

    filename = "MedBook_Development_Diary_20Days.docx"
    doc.save(filename)
    return filename

if __name__ == "__main__":
    fn = create_diary()
    print(f"Nhật ký đã được tạo: {fn}")
