
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('BÁO CÁO TIẾN ĐỘ DỰ ÁN MEDBOOK', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    p = doc.add_paragraph()
    p.add_run('Người thực hiện: ').bold = True
    p.add_run('Trương Thế Hải Thịnh & Nguyễn Thị Quỳnh Trang\n')
    p.add_run('Ngày báo cáo: ').bold = True
    p.add_run('13/04/2026\n')
    p.add_run('Giai đoạn: ').bold = True
    p.add_run('Hoàn tất MVP (Minimum Viable Product)')

    # Section 1: Major Changes
    doc.add_heading('1. CÁC THAY ĐỔI VÀ CẢI TIẾN SO VỚI PROPOSAL GỐC', level=1)
    
    changes = [
        ("Hệ thống Đơn thuốc chuyên nghiệp", "Bổ sung module kê đơn thuốc đầy đủ với các trường dữ liệu y tế chuẩn: Chẩn đoán, Lời dặn, Chiều cao, Cân nặng."),
        ("Chuẩn hóa định dạng liều dùng", "Áp dụng định dạng 'Sáng: 1 viên   Trưa: 0...' giúp bệnh nhân dễ đọc và bác sĩ thao tác nhanh."),
        ("Quy trình In & Xem hồ sơ A5", "Thiết kế template đơn thuốc khổ A5 chuẩn quy định. Bác sĩ có quyền in, bệnh nhân chỉ có quyền xem bản sao điện tử."),
        ("Hệ thống biểu tượng chuyên khoa", "Thay thế icon stethoscope chung bằng bộ icon đặc thù cho từng khoa (Tim mạch, Thần kinh, Mắt, Da liễu...)"),
        ("Chỉnh sửa ca khám (Edit Mode)", "Cho phép bác sĩ sửa lại hồ sơ/đơn thuốc sau khi đã hoàn thành ca khám để khắc phục sai sót."),
        ("Đồng bộ Thương hiệu (Branding)", "Thay thế toàn bộ các thông tin cũ (BV 175) thành 'BV MedBook' trên toàn hệ thống.")
    ]

    for title, desc in changes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # Section 2: Progress Status
    doc.add_heading('2. TÌNH TRẠNG TIẾN ĐỘ MVP (DONE)', level=1)
    status_table = doc.add_table(rows=1, cols=3)
    status_table.style = 'Table Grid'
    hdr_cells = status_table.rows[0].cells
    hdr_cells[0].text = 'Module'
    hdr_cells[1].text = 'Tính năng'
    hdr_cells[2].text = 'Trạng thái'

    data = [
        ("Auth", "JWT, Login/Register 3 Roles", "100%"),
        ("Specialty", "Admin CRUD, Specialty Icons", "100%"),
        ("Doctor", "Schedules, Working Hours", "100%"),
        ("Appointment", "Booking, Confirm, Cancel", "100%"),
        ("Prescription", "A5 Print, Electronic Record", "100%"),
        ("UI/UX", "Responsive, Professional Icons", "100%"),
    ]

    for mod, feat, stat in data:
        row_cells = status_table.add_row().cells
        row_cells[0].text = mod
        row_cells[1].text = feat
        row_cells[2].text = stat

    # Section 3: Next Steps
    doc.add_heading('3. KẾ HOẠCH CHO GIAI ĐOẠN BETA (NGÀY MAI)', level=1)
    doc.add_paragraph("Dựa trên tiến độ hiện tại, các bước tiếp theo cần thực hiện trong ngày mai:")
    
    next_steps = [
        "Email nhắc lịch tự động: Cấu hình APScheduler để gửi email nhắc nhở 24h trước ca khám.",
        "Tính năng Block Date: Cho phép bác sĩ đánh dấu các ngày nghỉ lễ/nghỉ đột xuất.",
        "Hệ thống Đánh giá (Review): Cho phép bệnh nhân gửi Feedback sau khi khám xong.",
        "Dashboard Thống kê (Charts): Vẽ biểu đồ tăng trưởng lịch hẹn cho Admin."
    ]
    for step in next_steps:
        doc.add_paragraph(step, style='List Number')

    # Save
    filename = "MedBook_Progress_Report_MVP.docx"
    doc.save(filename)
    return filename

if __name__ == "__main__":
    fn = create_report()
    print(f"Báo cáo đã được tạo: {fn}")
